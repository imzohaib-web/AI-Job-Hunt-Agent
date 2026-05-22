"""
agents/resume_tailor.py — AGENT 5: Tailor resume for a target job (LLM + DOCX).
"""

import json
import logging
import re
from pathlib import Path
from typing import Dict, Optional

logger = logging.getLogger("job_agent.resume_tailor")

TAILOR_SYSTEM = """You are an expert resume writer and ATS optimization specialist.
Return ONLY valid JSON. Never fabricate experience."""

TAILOR_PROMPT = """
Tailor this candidate's resume for the specific job below.

CANDIDATE PROFILE:
{profile_json}

TARGET JOB:
Title: {job_title}
Company: {company}
Description: {job_description}

Return JSON:
{{
  "tailored_summary": "2-3 sentence summary for this role",
  "skills_to_highlight": ["skill1", "skill2"],
  "experience": [
    {{
      "title": "Job Title",
      "company": "Company",
      "duration": "dates",
      "bullets": ["achievement with keywords"]
    }}
  ],
  "ats_keywords_added": ["keyword1"],
  "tailoring_notes": "Brief explanation of changes"
}}
"""


def tailor_resume_with_llm(profile: Dict, job: Dict) -> Dict:
    from core.llm_client import call_llm_json

    profile_subset = {
        "name": profile.get("name", ""),
        "summary": profile.get("summary", "")[:300],
        "skills": profile.get("skills", []),
        "experience": profile.get("experience", [])[:4],
        "projects": profile.get("projects", [])[:3],
        "education": profile.get("education", []),
    }
    prompt = TAILOR_PROMPT.format(
        profile_json=json.dumps(profile_subset, indent=2),
        job_title=job.get("title", ""),
        company=job.get("company", ""),
        job_description=(job.get("description") or "")[:2000],
    )
    result = call_llm_json(prompt, system=TAILOR_SYSTEM)
    if "error" in result:
        result = {
            "tailored_summary": profile.get("summary", ""),
            "skills_to_highlight": profile.get("skills", [])[:12],
            "experience": profile.get("experience", []),
            "ats_keywords_added": [],
            "tailoring_notes": "LLM tailoring unavailable — using original profile.",
        }
    return result


def build_docx_resume(profile: Dict, tailored: Dict, job: Dict, output_dir: Path) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    name = profile.get("name", "Candidate")
    title = doc.add_paragraph()
    title.add_run(name).bold = True
    title.runs[0].font.size = Pt(16)

    contact = doc.add_paragraph()
    contact.add_run(
        " | ".join(filter(None, [profile.get("email", ""), profile.get("phone", "")]))
    )
    contact.runs[0].font.size = Pt(9)

    summary = tailored.get("tailored_summary") or profile.get("summary", "")
    if summary:
        doc.add_paragraph("PROFESSIONAL SUMMARY").runs[0].bold = True
        doc.add_paragraph(summary)

    skills = tailored.get("skills_to_highlight") or profile.get("skills", [])
    if skills:
        doc.add_paragraph("SKILLS").runs[0].bold = True
        doc.add_paragraph("  •  ".join(skills[:12]))

    experience = tailored.get("experience") or profile.get("experience", [])
    if experience:
        doc.add_paragraph("EXPERIENCE").runs[0].bold = True
        for exp in experience:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('title', '')} | {exp.get('company', '')}").bold = True
            for bullet in (exp.get("bullets") or exp.get("achievements") or [])[:5]:
                doc.add_paragraph(f"• {bullet}", style="List Bullet")

    safe_company = re.sub(r"[^\w\s-]", "", job.get("company", "company")).replace(" ", "_")
    safe_title = re.sub(r"[^\w\s-]", "", job.get("title", "role")).replace(" ", "_")
    filename = f"Resume_{name.replace(' ', '_')}_{safe_company}_{safe_title}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    return str(filepath)


def run_resume_tailor(profile: Dict, job: Dict, output_dir: Optional[Path] = None) -> Dict:
    from core.config import OUTPUT_DIR

    if output_dir is None:
        output_dir = OUTPUT_DIR
    output_dir.mkdir(exist_ok=True)

    logger.info("=== Resume Tailor Agent START ===")
    tailored = tailor_resume_with_llm(profile, job)
    resume_path = build_docx_resume(profile, tailored, job, output_dir)
    tailored["resume_path"] = resume_path
    logger.info("=== Resume Tailor Agent DONE ===")
    return tailored
