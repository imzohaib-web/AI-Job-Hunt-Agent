"""
agents/cover_letter.py — AGENT 6: Generate personalized cover letter.
"""

import logging
import re
from pathlib import Path
from typing import Dict, Optional
from datetime import date

logger = logging.getLogger("job_agent.cover_letter")

COVER_LETTER_SYSTEM = """You are an expert career coach who writes compelling cover letters.
3-4 paragraphs, ~300 words. Write the letter directly — no meta-commentary."""

COVER_LETTER_PROMPT = """
Write a personalized cover letter.

CANDIDATE:
Name: {name}
Skills: {skills}
Experience: {experience_summary}
Summary: {summary}

JOB:
Title: {job_title}
Company: {company}
{company_context}
Description: {job_description}

Write the complete cover letter now:
"""


def generate_cover_letter(profile: Dict, job: Dict, company_info: Optional[Dict] = None) -> str:
    from core.llm_client import call_llm

    exp_parts = []
    for exp in profile.get("experience", [])[:2]:
        achievements = exp.get("achievements") or exp.get("bullets", [])
        bullet = achievements[0] if achievements else ""
        exp_parts.append(f"{exp.get('title')} at {exp.get('company')}: {bullet}")

    company_context = ""
    if company_info:
        company_context = (
            f"Industry: {company_info.get('industry', '')}\n"
            f"Mission: {company_info.get('mission', '')}\n"
            f"Why apply: {company_info.get('why_apply', '')}"
        )

    prompt = COVER_LETTER_PROMPT.format(
        name=profile.get("name", "Candidate"),
        skills=", ".join(profile.get("skills", [])[:12]),
        experience_summary="\n".join(exp_parts) or "Relevant student / junior experience",
        summary=profile.get("summary", "")[:400],
        job_title=job.get("title", ""),
        company=job.get("company", ""),
        company_context=company_context,
        job_description=(job.get("description") or "")[:1500],
    )
    return call_llm(prompt, system=COVER_LETTER_SYSTEM)


def save_cover_letter_txt(content: str, profile: Dict, job: Dict, output_dir: Path) -> str:
    safe = re.sub(r"[^\w\s-]", "", job.get("company", "co")).replace(" ", "_")
    path = output_dir / f"CoverLetter_{profile.get('name', 'User').replace(' ', '_')}_{safe}.txt"
    path.write_text(content, encoding="utf-8")
    return str(path)


def save_cover_letter_docx(content: str, profile: Dict, job: Dict, output_dir: Path) -> str:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph(f"Dear Hiring Manager at {job.get('company', '')},")
    for para in content.strip().split("\n\n"):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            if p.runs:
                p.runs[0].font.size = Pt(11)
    safe = re.sub(r"[^\w\s-]", "", job.get("company", "co")).replace(" ", "_")
    path = output_dir / f"CoverLetter_{profile.get('name', 'User').replace(' ', '_')}_{safe}.docx"
    doc.save(str(path))
    return str(path)


def run_cover_letter_agent(
    profile: Dict,
    job: Dict,
    company_info: Optional[Dict] = None,
    output_dir: Optional[Path] = None,
) -> Dict:
    from core.config import OUTPUT_DIR

    if output_dir is None:
        output_dir = OUTPUT_DIR
    output_dir.mkdir(exist_ok=True)

    logger.info("=== Cover Letter Agent START ===")
    content = generate_cover_letter(profile, job, company_info)
    txt_path = save_cover_letter_txt(content, profile, job, output_dir)
    docx_path = save_cover_letter_docx(content, profile, job, output_dir)
    logger.info("=== Cover Letter Agent DONE ===")
    return {"content": content, "txt_path": txt_path, "docx_path": docx_path}
