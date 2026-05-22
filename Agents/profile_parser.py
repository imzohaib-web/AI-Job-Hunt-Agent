"""
agents/profile_parser.py — AGENT 1: Parse resume PDF/DOCX into structured profile.
"""

import json
import logging
import re
from pathlib import Path
from typing import Dict, Optional

logger = logging.getLogger("job_agent.profile_parser")

PROFILE_SYSTEM = """You are an expert resume parser. Extract structured candidate data.
Return ONLY valid JSON. Never invent employers or degrees not present in the text."""

PROFILE_PROMPT = """
Extract a structured profile from this resume text.

Return JSON:
{{
  "name": "Full name",
  "email": "email or empty",
  "phone": "phone or empty",
  "summary": "2-3 sentence professional summary",
  "skills": ["skill1", "skill2"],
  "experience": [
    {{
      "title": "Role",
      "company": "Company",
      "duration": "dates",
      "achievements": ["bullet1", "bullet2"]
    }}
  ],
  "education": [{{"degree": "", "institution": "", "year": "", "gpa": ""}}],
  "projects": [{{"name": "", "description": "", "technologies": [], "impact": ""}}],
  "certifications": [{{"name": "", "issuer": "", "year": ""}}]
}}

RESUME TEXT:
{raw_text}
"""


def extract_text_from_pdf(file_path: str) -> str:
    from pdfminer.high_level import extract_text

    text = extract_text(file_path)
    logger.info("Extracted %d chars from PDF: %s", len(text), file_path)
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    text = "\n".join(paragraphs)
    logger.info("Extracted %d chars from DOCX: %s", len(text), file_path)
    return text.strip()


def extract_resume_text(file_path: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    if suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported resume format: {suffix}")


def parse_profile_with_llm(raw_text: str) -> Dict:
    from core.llm_client import call_llm_json

    trimmed = raw_text[:8000]
    prompt = PROFILE_PROMPT.format(raw_text=trimmed)
    result = call_llm_json(prompt, system=PROFILE_SYSTEM)

    if "error" in result:
        logger.warning("LLM profile parse failed, using minimal profile.")
        return {
            "name": "Candidate",
            "email": "",
            "skills": [],
            "experience": [],
            "education": [],
            "projects": [],
            "certifications": [],
            "summary": trimmed[:500],
        }
    return result


def save_profile_to_db(profile: Dict, raw_text: str) -> int:
    from core.config import get_db

    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO profiles (name, email, raw_text, skills, experience, education, certifications)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            profile.get("name", ""),
            profile.get("email", ""),
            raw_text[:20000],
            json.dumps(profile.get("skills", [])),
            json.dumps(profile.get("experience", [])),
            json.dumps(profile.get("education", [])),
            json.dumps(profile.get("certifications", [])),
        ),
    )
    profile_id = cur.lastrowid
    conn.commit()
    conn.close()
    return profile_id


def run_profile_parser(resume_path: str) -> Dict:
    """Parse resume file and persist profile to SQLite."""
    logger.info("=== Profile Parser START ===")
    raw_text = extract_resume_text(resume_path)
    if len(raw_text) < 50:
        raise ValueError("Resume text too short — check file quality.")

    profile = parse_profile_with_llm(raw_text)
    profile["raw_text"] = raw_text
    profile["profile_id"] = save_profile_to_db(profile, raw_text)
    logger.info("=== Profile Parser DONE — %s ===", profile.get("name"))
    return profile
